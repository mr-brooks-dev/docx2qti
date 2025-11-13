#!/usr/bin/env python3
"""
highlight_quiz.py
Robust version: reads answers from answers.txt, highlights matches in a .docx.
Usage examples:
  python highlight_quiz.py
  python highlight_quiz.py -i "MyQuiz.docx" -o "MyQuiz_highlighted.docx" -a answers.txt
"""

import argparse
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import sys

DEFAULT_OUTPUT_SUFFIX = "_Highlighted_Answers.docx"
DEFAULT_ANSWERS_FILE = "answers.txt"

def load_answers(file_path: Path):
    if not file_path.exists():
        raise FileNotFoundError(f"Answers file not found: {file_path}")
    text = file_path.read_text(encoding="utf-8")
    # Split on one or more blank lines, preserving multi-line answers as single block
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    return blocks

def normalize(text: str) -> str:
    text = text or ""
    text = text.lower()
    text = text.replace("‘", "'").replace("’", "'").replace("“", '"').replace("”", '"')
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def highlight_paragraph(paragraph, color=WD_COLOR_INDEX.GREEN):
    for run in paragraph.runs:
        run.font.highlight_color = color

def paragraph_matches(text: str, normalized_answers):
    norm_text = normalize(text)
    return any(ans in norm_text for ans in normalized_answers)

def find_docx_files(search_dir: Path):
    # ignore temporary files (starting with ~) and the generated output if present
    return [p for p in sorted(search_dir.glob("*.docx")) if not p.name.startswith("~")]

def choose_from_list(items):
    for idx, item in enumerate(items, start=1):
        print(f"  [{idx}] {item.name}")
    while True:
        choice = input(f"Select file number (1-{len(items)}) or 'q' to quit: ").strip()
        if choice.lower() == 'q':
            sys.exit("Aborted by user.")
        if choice.isdigit():
            i = int(choice)
            if 1 <= i <= len(items):
                return items[i-1]
        print("Invalid choice — try again.")

def highlight_matches_in_docx(input_path: Path, output_path: Path, answers):
    try:
        doc = Document(str(input_path))
    except Exception as e:
        raise RuntimeError(f"Failed to open Word file '{input_path}': {e}")

    normalized_answers = [normalize(a) for a in answers]

    # highlight in body paragraphs
    for para in doc.paragraphs:
        if paragraph_matches(para.text, normalized_answers):
            highlight_paragraph(para)

    # highlight in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if paragraph_matches(para.text, normalized_answers):
                        highlight_paragraph(para)

    doc.save(str(output_path))
    print(f"✅ Highlighted answers saved to: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Highlight quiz answers in a .docx using answers.txt")
    parser.add_argument("-i", "--input", help="Input .docx file (optional). If omitted, script searches current folder.")
    parser.add_argument("-o", "--output", help="Output .docx file (optional). Defaults to input filename + suffix.")
    parser.add_argument("-a", "--answers", default=DEFAULT_ANSWERS_FILE, help=f"Answers file (default: {DEFAULT_ANSWERS_FILE})")
    args = parser.parse_args()

    cwd = Path.cwd()

    # Resolve answers file
    answers_path = Path(args.answers)
    if not answers_path.is_absolute():
        answers_path = cwd / answers_path

    if not answers_path.exists():
        print(f"ERROR: answers file not found: {answers_path}")
        print("Create an answers.txt in this folder, or pass -a <path_to_answers.txt>")
        sys.exit(1)

    # Resolve input docx
    if args.input:
        input_path = Path(args.input)
        if not input_path.is_absolute():
            input_path = cwd / input_path
        if not input_path.exists():
            print(f"ERROR: Specified input file not found: {input_path}")
            sys.exit(1)
    else:
        docx_files = find_docx_files(cwd)
        if not docx_files:
            print("ERROR: No .docx files found in the current directory.")
            print("Place your quiz .docx in the same folder as this script or run with -i <file.docx>")
            sys.exit(1)
        if len(docx_files) == 1:
            input_path = docx_files[0]
            print(f"Auto-detected input file: {input_path.name}")
        else:
            print("Multiple .docx files found in the current directory:")
            input_path = choose_from_list(docx_files)
            print(f"Using: {input_path.name}")

    # Resolve output path
    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = cwd / output_path
    else:
        output_path = input_path.with_name(input_path.stem + DEFAULT_OUTPUT_SUFFIX)

    # Load answers and run
    try:
        answers = load_answers(answers_path)
        if not answers:
            print(f"ERROR: No answers parsed from {answers_path}. Make sure the file has content and blocks separated by blank lines.")
            sys.exit(1)
    except Exception as e:
        print(f"ERROR reading answers: {e}")
        sys.exit(1)

    try:
        highlight_matches_in_docx(input_path, output_path, answers)
    except Exception as e:
        print(f"ERROR during processing: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
